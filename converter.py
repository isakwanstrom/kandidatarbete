from docx import Document
from lxml import etree
from pictureExtractor import extract_images_from_docx

def is_bullet_paragraph(paragraph):
    return paragraph.style.name.startswith('List Bullet')

def is_numbered_paragraph(paragraph):
    return paragraph.style.name.startswith('List Number')

# Ladda Word-dokumentet
document = Document(r'C:\Users\wante\Downloads\Regelbok (prototyp 1) (2).docx')

# Skapa XML-strukturen
root = etree.Element('regelbok')
current_rubrik = None
current_underrubrik = None
current_subunderrubrik = None
previous_paragraph_style = None
current_list = None

# Kör bildextraktionsfunktionen för att få bildpositioner
image_positions, image_src = extract_images_from_docx(r'C:\Users\wante\Downloads\Regelbok (prototyp 1) (2).docx', r"C:\Users\wante\Desktop\Kandidatarbete\hemsida\bilder")


for paragraph in document.paragraphs:
    # Kontrollerar om paragrafen är en del av en lista
    if is_bullet_paragraph(paragraph) or is_numbered_paragraph(paragraph):
        # Kontrollera vilken typ av lista det är och skapa en
        list_type = 'ul' if is_bullet_paragraph(paragraph) else 'ol'
        if current_list is None or current_list.tag != list_type:
            current_list = etree.SubElement(root, list_type)
        # Lägg till listelementet
        list_item = etree.SubElement(current_list, 'li')
        list_item.text = paragraph.text.strip()
    else:
        # Inte en lista, eller slutet på nuvarande lista
        current_list = None
        # Här lägger du till logiken för att skapa 'info'-element osv.


# Förbered en behållare för element som kan innehålla bilder
element_for_images = {}
# Struktur för att hålla text och bildpositioner
paragraph_content = []

for i, paragraph in enumerate(document.paragraphs):
    current_element = None
    if paragraph.style.name == 'Heading 1' and paragraph.text.strip() != "":
        current_rubrik = etree.SubElement(root, 'rubrik', namn=paragraph.text.strip(), klass='heading1')
        current_underrubrik = None
        current_subunderrubrik = None
        current_element = current_rubrik
    elif paragraph.style.name == 'Heading 2' and current_rubrik is not None and paragraph.text.strip() != "":
        current_underrubrik = etree.SubElement(current_rubrik, 'underrubrik', namn=paragraph.text.strip(), klass='heading2')
        current_subunderrubrik = None
        current_element = current_underrubrik
    elif paragraph.style.name == 'Heading 3' and paragraph.text.strip() != "" and current_underrubrik is not None:
        current_subunderrubrik = etree.SubElement(current_underrubrik, 'underrubrik', namn=paragraph.text.strip(), klass='heading3')
        current_element = current_subunderrubrik
    elif paragraph.text.strip():
        if current_subunderrubrik is not None:
            text = etree.SubElement(current_subunderrubrik, 'text', klass='info')
        elif current_underrubrik is not None:
            text = etree.SubElement(current_underrubrik, 'text', klass='info')
        elif current_rubrik is not None:
            text = etree.SubElement(current_rubrik, 'text', klass='rubriktext')
        else:
            text = etree.SubElement(root, 'text', klass='info')
            runs_content = []
            for run_index, run in enumerate(paragraph.runs):
                # Lägg till texten från varje run
                runs_content.append({"type": "text", "content": run.text})

                # Lägg till bilder som hör till denna run om det finns några
                images_for_run = [img for img in image_positions if img['paragraph_index'] == i and img['run_index'] == run_index]
                for img in images_for_run:
                    runs_content.append({"type": "image", "src": f"bilder/{img['filename']}"})

            # Spara den kombinerade innehållet för varje paragraf
            paragraph_content.append(runs_content)
        text.text = paragraph.text.strip()
        current_element = text

    # Map current element to its position for image placement
    if current_element is not None:
        element_for_images[i] = current_element

# Place images in their correct positions
for image_info in image_positions:
    paragraph_index = image_info['paragraph_index']
    if paragraph_index in element_for_images:
        img_tag = etree.SubElement(element_for_images[paragraph_index], 'img', src=f"bilder/{image_info['filename']}")

# Lägg till en XML Stylesheet Processing Instruction
pi = etree.ProcessingInstruction('xml-stylesheet', 'type="text/xsl" href="transformation.xsl"')
root.addprevious(pi)

# Spara den formaterade XML-strukturen till en fil
tree = etree.ElementTree(root)
tree.write('ut_export.xml', pretty_print=True, xml_declaration=True, encoding='UTF-8')
